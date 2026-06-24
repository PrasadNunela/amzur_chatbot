"""Business logic for multi-modal contract analysis."""

from __future__ import annotations

import asyncio
import io
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from uuid import UUID, uuid4

from fastapi import HTTPException
from openai import OpenAIError
from pypdf import PdfReader
from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.llm import openai_client
from app.core.config import settings
from app.core.logger import logger
from app.models.chat import ContractAnalysisReport
from app.schemas.contract_analysis import (
    ContractAnalysisResponseSchema,
    ContractClauseSchema,
    ContractDataExtractionSchema,
    ContractRiskSchema,
    ContractSummarySchema,
)

MAX_CONTRACT_TEXT_CHARS = 120000
CONTRACT_ANALYSIS_TIMEOUT_SECONDS = 180
CONTRACT_ANALYSIS_FALLBACK_TEXT_CHARS = 30000


class ContractAnalysisService:
    """Service for extracting and analyzing legal contracts."""

    @staticmethod
    def _to_optional_str(value: object) -> str | None:
        """Convert mixed scalar values to optional strings for schema safety."""
        if value is None:
            return None
        if isinstance(value, str):
            cleaned = value.strip()
            return cleaned if cleaned else None
        # LLMs may return numbers/bools for fields like contract_value.
        return str(value)

    @staticmethod
    def _normalize_extracted_data(payload: object) -> dict[str, object]:
        """Normalize extracted_data payload into schema-compatible primitives."""
        if not isinstance(payload, dict):
            payload = {}

        party_names_raw = payload.get("party_names")
        party_names: list[str] = []
        if isinstance(party_names_raw, list):
            party_names = [str(item).strip() for item in party_names_raw if str(item).strip()]
        elif party_names_raw is not None:
            maybe_name = str(party_names_raw).strip()
            if maybe_name:
                party_names = [maybe_name]

        return {
            "party_names": party_names,
            "effective_date": ContractAnalysisService._to_optional_str(payload.get("effective_date")),
            "expiration_date": ContractAnalysisService._to_optional_str(payload.get("expiration_date")),
            "governing_law": ContractAnalysisService._to_optional_str(payload.get("governing_law")),
            "contract_value": ContractAnalysisService._to_optional_str(payload.get("contract_value")),
            "renewal_terms": ContractAnalysisService._to_optional_str(payload.get("renewal_terms")),
            "payment_terms": ContractAnalysisService._to_optional_str(payload.get("payment_terms")),
            "notice_period": ContractAnalysisService._to_optional_str(payload.get("notice_period")),
        }

    @staticmethod
    async def save_report(
        db: AsyncSession,
        *,
        user_id: UUID,
        report: ContractAnalysisResponseSchema,
        uploaded_file: tuple[str, str | None, bytes] | None = None,
    ) -> ContractAnalysisReport:
        """Persist a generated contract analysis report for a user."""
        uploaded_filename: str | None = None
        uploaded_file_path: str | None = None
        uploaded_file_mime_type: str | None = None
        uploaded_file_size: str | None = None

        if uploaded_file is not None:
            original_filename, mime_type, file_bytes = uploaded_file
            uploads_dir = Path(settings.UPLOAD_DIR) / "contract_reports"
            uploads_dir.mkdir(parents=True, exist_ok=True)

            safe_name = Path(original_filename).name
            stored_name = f"{user_id}_{uuid4()}_{safe_name}"
            stored_path = uploads_dir / stored_name
            stored_path.write_bytes(file_bytes)

            uploaded_filename = safe_name
            uploaded_file_path = str(stored_path.resolve())
            uploaded_file_mime_type = mime_type
            uploaded_file_size = str(len(file_bytes))

        record = ContractAnalysisReport(
            user_id=user_id,
            filename=report.filename,
            uploaded_filename=uploaded_filename,
            uploaded_file_path=uploaded_file_path,
            uploaded_file_mime_type=uploaded_file_mime_type,
            uploaded_file_size=uploaded_file_size,
            report_json=json.dumps(report.model_dump(mode="json")),
        )
        db.add(record)
        await db.commit()
        await db.refresh(record)
        return record

    @staticmethod
    async def list_reports(db: AsyncSession, *, user_id: UUID) -> list[ContractAnalysisReport]:
        """List saved reports for a user, newest first."""
        stmt = (
            select(ContractAnalysisReport)
            .where(ContractAnalysisReport.user_id == user_id)
            .order_by(ContractAnalysisReport.created_at.desc())
        )
        result = await db.execute(stmt)
        return result.scalars().all()

    @staticmethod
    async def get_report(
        db: AsyncSession,
        *,
        report_id: UUID,
        user_id: UUID,
    ) -> ContractAnalysisReport | None:
        """Get one saved report owned by a user."""
        stmt = select(ContractAnalysisReport).where(
            ContractAnalysisReport.id == report_id,
            ContractAnalysisReport.user_id == user_id,
        )
        result = await db.execute(stmt)
        return result.scalars().first()

    @staticmethod
    async def delete_report(
        db: AsyncSession,
        *,
        report_id: UUID,
        user_id: UUID,
    ) -> bool:
        """Delete a saved report if it belongs to the user."""
        report = await ContractAnalysisService.get_report(db, report_id=report_id, user_id=user_id)
        if not report:
            return False

        if report.uploaded_file_path:
            try:
                file_path = Path(report.uploaded_file_path)
                if file_path.exists():
                    file_path.unlink()
            except Exception:
                # Keep delete resilient even if file cleanup fails.
                pass

        await db.delete(report)
        await db.commit()
        return True

    @staticmethod
    def get_uploaded_file_path(record: ContractAnalysisReport) -> Path | None:
        """Get uploaded contract file path if available and existing."""
        if not record.uploaded_file_path:
            return None
        path = Path(record.uploaded_file_path)
        if not path.exists():
            return None
        return path

    @staticmethod
    def deserialize_report(record: ContractAnalysisReport) -> ContractAnalysisResponseSchema:
        """Deserialize stored JSON report into schema."""
        try:
            payload = json.loads(record.report_json)
        except json.JSONDecodeError as exc:
            raise HTTPException(
                status_code=500,
                detail={"error": "invalid_saved_report", "message": "Saved report payload is corrupted."},
            ) from exc
        return ContractAnalysisResponseSchema.model_validate(payload)

    @staticmethod
    def extract_contract_text(filename: str, file_bytes: bytes) -> str:
        """Extract plain text from supported contract formats."""
        lower = filename.lower()
        if lower.endswith(".pdf"):
            return ContractAnalysisService._extract_pdf_text(file_bytes)
        if lower.endswith(".docx"):
            return ContractAnalysisService._extract_docx_text(file_bytes)
        raise HTTPException(
            status_code=400,
            detail={"error": "invalid_file_type", "message": "Only PDF and DOCX files are supported."},
        )

    @staticmethod
    def _extract_pdf_text(file_bytes: bytes) -> str:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        return "\n\n".join(pages).strip()

    @staticmethod
    def _extract_docx_text(file_bytes: bytes) -> str:
        doc = Document(io.BytesIO(file_bytes))
        chunks: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    chunks.append(" | ".join(cells))

        return "\n".join(chunks).strip()

    @staticmethod
    def _truncate_contract_text(contract_text: str) -> str:
        if len(contract_text) <= MAX_CONTRACT_TEXT_CHARS:
            return contract_text
        return (
            contract_text[:MAX_CONTRACT_TEXT_CHARS]
            + "\n\n[Content truncated for analysis token safety]"
        )

    @staticmethod
    def _build_messages(system_prompt: str, user_prompt: str) -> list[dict[str, str]]:
        return [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

    @staticmethod
    def _pick_fallback_model(primary_model: str) -> str:
        """Choose a secondary model for resilience when the primary call fails."""
        if primary_model == "gpt-4o":
            return "gemini/gemini-2.5-flash"
        return "gpt-4o"

    @staticmethod
    def _build_fallback_contract_text(contract_text: str) -> str:
        """Trim contract text aggressively for fallback analysis attempts."""
        if len(contract_text) <= CONTRACT_ANALYSIS_FALLBACK_TEXT_CHARS:
            return contract_text
        return (
            contract_text[:CONTRACT_ANALYSIS_FALLBACK_TEXT_CHARS]
            + "\n\n[Content reduced for fallback analysis due upstream timeout]"
        )

    @staticmethod
    def _build_degraded_response(
        *,
        filename: str,
        contract_text: str,
        reason: str,
    ) -> ContractAnalysisResponseSchema:
        """Return a safe best-effort report when upstream LLM calls are unavailable."""
        snippet = contract_text[:1200].strip()
        summary_text = (
            "Best-effort analysis generated because upstream AI provider was unavailable "
            f"({reason}). This report includes limited deterministic extraction only."
        )
        if snippet:
            summary_text += f"\n\nPreview:\n{snippet}"

        extracted_data = ContractDataExtractionSchema.model_validate(
            {
                "party_names": [],
                "effective_date": None,
                "expiration_date": None,
                "governing_law": None,
                "contract_value": None,
                "renewal_terms": None,
                "payment_terms": None,
                "notice_period": None,
            }
        )

        return ContractAnalysisResponseSchema(
            filename=filename,
            summary=ContractSummarySchema.model_validate(
                {
                    "executive_summary": summary_text,
                    "key_terms": ["degraded_mode", "llm_unavailable"],
                }
            ),
            clauses=[],
            risks=[],
            extracted_data=extracted_data,
            analyzed_at=datetime.now(timezone.utc),
        )

    @staticmethod
    def _safe_json_load(text: str) -> dict:
        cleaned = text.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
            cleaned = re.sub(r"```$", "", cleaned).strip()

        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            match = re.search(r"\{.*\}", cleaned, re.DOTALL)
            if match:
                return json.loads(match.group(0))
            raise

    @staticmethod
    def _run_json_task(
        *,
        task_name: str,
        task_prompt: str,
        contract_text: str | None,
        user_email: str,
        model: str,
    ) -> dict:
        client = openai_client()
        system_prompt = (
            "You are a legal contract analysis assistant. "
            "Return strictly valid JSON with no markdown or extra prose."
        )
        user_prompt = (
            f"Task: {task_name}\n\n"
            f"Instructions:\n{task_prompt}"
        )
        if contract_text is not None and contract_text.strip():
            user_prompt += f"\n\nContract:\n{contract_text}"

        messages = ContractAnalysisService._build_messages(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
        )

        response = client.chat.completions.create(
            model=model,
            messages=messages,
            user=user_email,
            temperature=0.1,
            extra_body={
                "metadata": {
                    "application": settings.APP_NAME,
                    "environment": settings.ENVIRONMENT,
                }
            },
        )
        content = response.choices[0].message.content or "{}"
        return ContractAnalysisService._safe_json_load(content)

    @staticmethod
    async def analyze_contract(
        *,
        filename: str,
        file_bytes: bytes,
        user_email: str,
    ) -> ContractAnalysisResponseSchema:
        contract_text = ContractAnalysisService.extract_contract_text(filename, file_bytes)
        if not contract_text.strip():
            raise HTTPException(
                status_code=400,
                detail={
                    "error": "empty_contract_text",
                    "message": "No readable text could be extracted from the uploaded document.",
                },
            )

        contract_text = ContractAnalysisService._truncate_contract_text(contract_text)

        full_analysis_prompt = (
            "Perform a complete contract analysis and return a single JSON object with these keys: "
            "summary, clauses, risks, extracted_data.\n\n"
            "1) summary: object with fields executive_summary (string) and key_terms (array of strings).\n"
            "2) clauses: array of objects with fields category, clause_title, description, source_excerpt.\n"
            "3) risks: array of objects with fields title, severity (low|medium|high), description, "
            "clause_reference, recommendation.\n"
            "4) extracted_data: object with fields party_names (array), effective_date, expiration_date, "
            "governing_law, contract_value, renewal_terms, payment_terms, notice_period.\n\n"
            "Use null for unknown single-value extracted fields and [] for unknown arrays. "
            "Return valid JSON only, with no markdown."
        )

        primary_model = settings.LLM_MODEL
        fallback_model = ContractAnalysisService._pick_fallback_model(primary_model)
        try:
            analysis_result = await asyncio.wait_for(
                asyncio.to_thread(
                    ContractAnalysisService._run_json_task,
                    task_name="Complete Contract Analysis",
                    task_prompt=full_analysis_prompt,
                    contract_text=contract_text,
                    user_email=user_email,
                    model=primary_model,
                ),
                timeout=CONTRACT_ANALYSIS_TIMEOUT_SECONDS,
            )
        except asyncio.TimeoutError as exc:
            try:
                fallback_text = ContractAnalysisService._build_fallback_contract_text(contract_text)
                analysis_result = await asyncio.wait_for(
                    asyncio.to_thread(
                        ContractAnalysisService._run_json_task,
                        task_name="Complete Contract Analysis",
                        task_prompt=full_analysis_prompt,
                        contract_text=fallback_text,
                        user_email=user_email,
                        model=fallback_model,
                    ),
                    timeout=CONTRACT_ANALYSIS_TIMEOUT_SECONDS,
                )
            except Exception as fallback_exc:
                logger.warning(
                    "Contract analysis degraded fallback engaged after timeout failures: %s",
                    fallback_exc,
                )
                return ContractAnalysisService._build_degraded_response(
                    filename=filename,
                    contract_text=contract_text,
                    reason="analysis_timeout",
                )
        except OpenAIError as exc:
            try:
                fallback_text = ContractAnalysisService._build_fallback_contract_text(contract_text)
                analysis_result = await asyncio.wait_for(
                    asyncio.to_thread(
                        ContractAnalysisService._run_json_task,
                        task_name="Complete Contract Analysis",
                        task_prompt=full_analysis_prompt,
                        contract_text=fallback_text,
                        user_email=user_email,
                        model=fallback_model,
                    ),
                    timeout=CONTRACT_ANALYSIS_TIMEOUT_SECONDS,
                )
            except asyncio.TimeoutError as fallback_exc:
                logger.warning(
                    "Contract analysis degraded fallback engaged after OpenAI timeout: %s",
                    fallback_exc,
                )
                return ContractAnalysisService._build_degraded_response(
                    filename=filename,
                    contract_text=contract_text,
                    reason="analysis_timeout",
                )
            except OpenAIError as fallback_exc:
                logger.warning(
                    "Contract analysis degraded fallback engaged after OpenAI failures: primary=%s fallback=%s",
                    exc,
                    fallback_exc,
                )
                return ContractAnalysisService._build_degraded_response(
                    filename=filename,
                    contract_text=contract_text,
                    reason="llm_error",
                )
        except json.JSONDecodeError as exc:
            raise HTTPException(
                status_code=502,
                detail={"error": "invalid_llm_json", "message": f"Model returned invalid JSON: {exc}"},
            ) from exc
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail={"error": "unexpected", "message": str(exc)},
            ) from exc

        clause_items: list[object] = []
        risk_items: list[object] = []
        extraction_payload_raw: object = {}
        summary_payload: object = {}
        if isinstance(analysis_result, dict):
            raw_clauses = analysis_result.get("clauses", [])
            raw_risks = analysis_result.get("risks", [])
            clause_items = raw_clauses if isinstance(raw_clauses, list) else []
            risk_items = raw_risks if isinstance(raw_risks, list) else []
            extraction_payload_raw = analysis_result.get("extracted_data", {})
            summary_payload = analysis_result.get("summary", {})

        summary_exec = "Summary unavailable."
        summary_terms: list[str] = []
        if isinstance(summary_payload, dict):
            maybe_exec = summary_payload.get("executive_summary")
            if isinstance(maybe_exec, str) and maybe_exec.strip():
                summary_exec = maybe_exec
            maybe_terms = summary_payload.get("key_terms")
            if isinstance(maybe_terms, list):
                summary_terms = [str(term).strip() for term in maybe_terms if str(term).strip()]

        clauses = [
            ContractClauseSchema.model_validate(item)
            for item in clause_items
            if isinstance(item, dict)
        ]
        risks = [
            ContractRiskSchema.model_validate(item)
            for item in risk_items
            if isinstance(item, dict)
        ]
        extracted_data_payload = ContractAnalysisService._normalize_extracted_data(
            extraction_payload_raw
        )

        summary = ContractSummarySchema.model_validate(
            {
                "executive_summary": summary_exec,
                "key_terms": summary_terms,
            }
        )
        extracted_data = ContractDataExtractionSchema.model_validate(extracted_data_payload)

        return ContractAnalysisResponseSchema(
            filename=filename,
            summary=summary,
            clauses=clauses,
            risks=risks,
            extracted_data=extracted_data,
            analyzed_at=datetime.now(timezone.utc),
        )
